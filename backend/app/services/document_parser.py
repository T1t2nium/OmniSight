"""Document parser — extract raw text from PDF and DOCX files.

Uses pdfplumber for PDF text extraction and python-docx for Word documents.
Returns a unified ParsedDocument for downstream entity extraction.

Usage:
    from app.services.document_parser import DocumentParser

    doc = DocumentParser.parse(pdf_bytes, "job_description.pdf")
    print(doc.raw_text)
"""

import io
import logging
from pathlib import Path

from app.models.interview import ParsedDocument

logger = logging.getLogger(__name__)


class UnsupportedFormatError(ValueError):
    """Raised when the file format is not PDF or DOCX."""


class DocumentParser:
    """Extract raw text from PDF and DOCX files.

    All methods are static — no state, no async needed (pdfplumber and
    python-docx are synchronous libraries).
    """

    # ---- Public API ----

    @staticmethod
    def parse(data: bytes, filename: str) -> ParsedDocument:
        """Auto-detect format by file extension and extract text.

        Args:
            data: Raw file bytes.
            filename: Original filename used to detect format.

        Returns:
            ParsedDocument with raw_text, pages, and metadata.

        Raises:
            UnsupportedFormatError: If the format is not PDF or DOCX.
        """
        ext = Path(filename).suffix.lower()
        if ext == ".pdf":
            return DocumentParser._parse_pdf(data, filename)
        elif ext in (".docx", ".doc"):
            return DocumentParser._parse_docx(data, filename)
        else:
            raise UnsupportedFormatError(
                f"Unsupported file format: {ext!r}. "
                f"Only PDF (.pdf) and Word (.docx) are supported."
            )

    # ---- PDF ----

    @staticmethod
    def _parse_pdf(data: bytes, filename: str) -> ParsedDocument:
        """Extract text from each page of a PDF using pdfplumber."""
        import pdfplumber

        pages: list[str] = []
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    pages.append(text)
                else:
                    pages.append("")

        raw_text = "\n\n".join(pages)
        logger.info(
            "PDF parsed: %s — %d pages, %d chars",
            filename, len(pages), len(raw_text),
        )
        return ParsedDocument(
            raw_text=raw_text,
            pages=pages,
            metadata={
                "file_name": filename,
                "page_count": len(pages),
                "format": "pdf",
            },
        )

    # ---- DOCX ----

    @staticmethod
    def _parse_docx(data: bytes, filename: str) -> ParsedDocument:
        """Extract text from paragraphs and tables in a DOCX file."""
        from docx import Document

        doc = Document(io.BytesIO(data))

        # Extract paragraphs
        paragraphs: list[str] = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # Extract table text
        for table in doc.tables:
            for row in table.rows:
                row_texts: list[str] = []
                for cell in row.cells:
                    ct = cell.text.strip()
                    if ct:
                        row_texts.append(ct)
                if row_texts:
                    paragraphs.append(" | ".join(row_texts))

        raw_text = "\n".join(paragraphs)
        logger.info(
            "DOCX parsed: %s — %d paragraphs, %d chars",
            filename, len(paragraphs), len(raw_text),
        )
        return ParsedDocument(
            raw_text=raw_text,
            pages=[raw_text],  # DOCX has no page concept — single page
            metadata={
                "file_name": filename,
                "page_count": 1,
                "format": "docx",
                "paragraph_count": len(paragraphs),
            },
        )
