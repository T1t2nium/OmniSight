"""Unit tests for DocumentParser — PDF and DOCX text extraction."""

from __future__ import annotations

import io
import tempfile
from pathlib import Path

import pytest

from app.models.interview import ParsedDocument
from app.services.document_parser import DocumentParser, UnsupportedFormatError


# ---- Fixtures ----


@pytest.fixture
def sample_docx_bytes() -> bytes:
    """Generate a minimal .docx file with known content."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("张三")
    doc.add_paragraph("Python 工程师")
    doc.add_paragraph("技能：Python, FastAPI, React, Docker")
    doc.add_paragraph("工作经历：")
    doc.add_paragraph("2020.06 - 2023.12  某科技有限公司  Python 开发工程师")
    doc.add_paragraph("负责后端 API 开发和数据库设计")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


@pytest.fixture
def sample_pdf_bytes() -> bytes:
    """Generate a valid PDF file with known English text content.

    Uses English text because fpdf2's built-in Helvetica font doesn't
    support CJK characters without adding a Unicode font file.
    """
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    pdf.multi_cell(
        0, 8,
        "Senior Python Engineer\n\n"
        "Responsibilities:\n"
        "1. Design and develop backend service architecture\n"
        "2. Perform system optimization and code review\n\n"
        "Requirements:\n"
        "- Bachelor's degree or above in CS or related field\n"
        "- 3+ years of Python development experience\n"
        "- Proficient in FastAPI and Django frameworks\n"
        "- Familiar with Docker and Kubernetes deployment\n"
        "- Preferred: Experience with React and TypeScript\n"
        "- Preferred: AI/ML project experience\n",
    )

    buf = io.BytesIO()
    pdf.output(buf)
    buf.seek(0)
    return buf.read()


# ---- Tests ----


class TestDocumentParser:
    """Test PDF and DOCX parsing via DocumentParser.parse()."""

    def test_parse_docx(self, sample_docx_bytes):
        """DOCX parsing extracts paragraphs correctly."""
        doc = DocumentParser.parse(sample_docx_bytes, "resume.docx")

        assert isinstance(doc, ParsedDocument)
        assert "张三" in doc.raw_text
        assert "Python 工程师" in doc.raw_text
        assert "FastAPI" in doc.raw_text
        assert "Docker" in doc.raw_text
        assert doc.metadata["format"] == "docx"
        assert doc.metadata["file_name"] == "resume.docx"
        assert len(doc.pages) == 1

    def test_parse_pdf(self, sample_pdf_bytes):
        """PDF parsing extracts text from pages."""
        doc = DocumentParser.parse(sample_pdf_bytes, "job_description.pdf")

        assert isinstance(doc, ParsedDocument)
        assert "Python" in doc.raw_text
        assert "FastAPI" in doc.raw_text
        assert "Docker" in doc.raw_text
        assert "Kubernetes" in doc.raw_text
        assert doc.metadata["format"] == "pdf"
        assert doc.metadata["file_name"] == "job_description.pdf"

    def test_parse_docx_via_doc_ext(self, sample_docx_bytes):
        """.doc extension is also handled as DOCX."""
        doc = DocumentParser.parse(sample_docx_bytes, "resume.doc")
        assert doc.metadata["format"] == "docx"

    def test_unsupported_format(self):
        """Unsupported formats raise UnsupportedFormatError."""
        with pytest.raises(UnsupportedFormatError, match="Unsupported"):
            DocumentParser.parse(b"dummy content", "file.txt")

    def test_parse_empty_docx(self):
        """Empty DOCX returns empty ParsedDocument."""
        from docx import Document

        doc = Document()
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        result = DocumentParser.parse(buf.read(), "empty.docx")
        assert result.raw_text == ""
        assert len(result.pages) == 1
        assert result.metadata["format"] == "docx"

    def test_parse_docx_with_table(self):
        """DOCX table content is extracted."""
        from docx import Document

        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "技能"
        table.rows[0].cells[1].text = "熟练程度"
        table.rows[1].cells[0].text = "Python"
        table.rows[1].cells[1].text = "精通"

        doc.add_paragraph("额外段落内容")

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        result = DocumentParser.parse(buf.read(), "table.docx")
        assert "Python" in result.raw_text
        assert "精通" in result.raw_text
        assert "额外段落内容" in result.raw_text
